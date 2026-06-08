from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
ACCENT   = RGBColor(0x6C, 0x63, 0xFF)   # purple-blue
DARK     = RGBColor(0x11, 0x11, 0x22)
MID      = RGBColor(0x44, 0x44, 0x66)
LIGHT    = RGBColor(0x88, 0x88, 0xAA)

# ── Helper: set paragraph spacing ─────────────────────────────────────────────
def set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spg = OxmlElement('w:spacing')
    spg.set(qn('w:before'), str(before))
    spg.set(qn('w:after'),  str(after))
    if line:
        spg.set(qn('w:line'),     str(line))
        spg.set(qn('w:lineRule'), 'auto')
    pPr.append(spg)

# ── Helper: horizontal rule ───────────────────────────────────────────────────
def add_hrule(doc, color="9999CC"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    set_spacing(p, before=0, after=60)
    return p

# ── Helper: section heading ───────────────────────────────────────────────────
def add_section_heading(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    set_spacing(p, before=180, after=40)
    add_hrule(doc)

# ── Helper: two-column line ───────────────────────────────────────────────────
def add_two_col(doc, left, right, left_bold=False, right_color=None):
    p  = doc.add_paragraph()
    r1 = p.add_run(left)
    r1.bold            = left_bold
    r1.font.size       = Pt(10)
    r1.font.color.rgb  = DARK
    tab = p.add_run('\t')
    r2 = p.add_run(right)
    r2.font.size       = Pt(9)
    r2.font.color.rgb  = right_color or LIGHT
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9260')
    tabs.append(tab_el)
    pPr.append(tabs)
    set_spacing(p, before=0, after=40)
    return p

# ── Helper: remove all table borders ─────────────────────────────────────────
def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

# ═══════════════════════════════════════════════════════════════════════════════
#  HEADER — two-column table (text left, photo right)
# ═══════════════════════════════════════════════════════════════════════════════
PHOTO_PATH = r"c:\Users\quo\repos\SebbePortfolio2\public\me.jpg"

header_tbl = doc.add_table(rows=1, cols=2)
remove_table_borders(header_tbl)

# Column widths: left ~11 cm, right ~6 cm (total 17 cm = A4 - margins)
header_tbl.columns[0].width = Cm(11)
header_tbl.columns[1].width = Cm(6)

left_cell  = header_tbl.cell(0, 0)
right_cell = header_tbl.cell(0, 1)

# ── Left cell: name / title / contact ────────────────────────────────────────
left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

name_p = left_cell.paragraphs[0]
run = name_p.add_run("Sebastian Larsson")
run.bold           = True
run.font.size      = Pt(26)
run.font.color.rgb = DARK
set_spacing(name_p, before=0, after=20)

title_p = left_cell.add_paragraph()
run2 = title_p.add_run("Full-Stack Developer  ·  AI Enthusiast")
run2.font.size      = Pt(12)
run2.font.color.rgb = ACCENT
set_spacing(title_p, before=0, after=40)

contact_p = left_cell.add_paragraph()
for i, item in enumerate([
    "sebbelarsson9601@gmail.com",
    "linkedin.com/in/sebastian-larsson-b45803246",
    "github.com/ManaInfectedRP",
    "Halmstad, Sweden",
]):
    r = contact_p.add_run(item)
    r.font.size      = Pt(9)
    r.font.color.rgb = MID
    if i < 3:
        sep = contact_p.add_run("   |   ")
        sep.font.size      = Pt(9)
        sep.font.color.rgb = LIGHT
set_spacing(contact_p, before=0, after=0)

# ── Right cell: photo ─────────────────────────────────────────────────────────
right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
photo_p = right_cell.paragraphs[0]
photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
if os.path.exists(PHOTO_PATH):
    run_img = photo_p.add_run()
    run_img.add_picture(PHOTO_PATH, width=Cm(5.5))

# Separator after header table
add_hrule(doc, "6C63FF")

# ═══════════════════════════════════════════════════════════════════════════════
#  PROFILE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Profile")

profile_p = doc.add_paragraph(
    "I started my career as a mechanic and warehouse worker in the automotive industry. "
    "That background taught me how to break down problems, follow procedures, and ship results — "
    "skills I now apply to building software. Today I'm focused on full-stack development and AI. "
    "I enjoy building real things people can use: video conferencing apps, REST APIs, desktop games, "
    "and AI experiments. Actively looking for work."
)
profile_p.runs[0].font.size       = Pt(10)
profile_p.runs[0].font.color.rgb  = MID
set_spacing(profile_p, before=0, after=40, line=276)

# ═══════════════════════════════════════════════════════════════════════════════
#  EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Experience")

experience = [
    ("2026",        "2nd LIA Full-Stack Developer",              "Podmanager.AI"),
    ("2025",        "Full-Stack Developer LIA Intern",           "Podmanager.AI"),
    ("2024 – 2026", "Software Development Student (AI spec.)",   "NBI/Handelsakademin, Halmstad"),
    ("2023",        "Warehouse Worker & Picker",                 "Enyroom AB, Halmstad"),
    ("2019 – 2020", "Mechanic",                                  "Göstorps Nissan Bil AB, Laholm"),
]

for period, role, place in experience:
    p = doc.add_paragraph()
    r1 = p.add_run(role)
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = DARK
    r2 = p.add_run(f"  —  {place}")
    r2.font.size      = Pt(9)
    r2.font.color.rgb = MID
    # right-aligned period
    r3 = p.add_run(f"\t{period}")
    r3.font.size      = Pt(9)
    r3.font.color.rgb = LIGHT
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9260')
    tabs.append(tab_el)
    pPr.append(tabs)
    set_spacing(p, before=0, after=60)

# ═══════════════════════════════════════════════════════════════════════════════
#  TECHNICAL SKILLS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Technical Skills")

skill_groups = [
    ("Languages",   "TypeScript · JavaScript · Python · C# · HTML / CSS · SQL"),
    ("Frontend",    "React · React Native · Vite · Tailwind CSS"),
    ("Backend",     "Node.js · Express · FastAPI · .NET · REST APIs · WebRTC"),
    ("Databases",   "PostgreSQL · MongoDB · SQLite · Entity Framework Core"),
    ("AI / ML",     "PyTorch · OpenAI API · Coqui TTS · MusicGen · AI tooling"),
    ("DevOps",      "Docker · Git · Render · Nginx · CI/CD"),
    ("Other",       "Unity · Pygame · Tkinter · Expo (React Native) · Zustand"),
]

for label, items in skill_groups:
    p  = doc.add_paragraph()
    r1 = p.add_run(f"{label}:  ")
    r1.bold           = True
    r1.font.size      = Pt(9)
    r1.font.color.rgb = DARK
    r2 = p.add_run(items)
    r2.font.size      = Pt(9)
    r2.font.color.rgb = MID
    set_spacing(p, before=0, after=30)

# ═══════════════════════════════════════════════════════════════════════════════
#  SELECTED PROJECTS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Selected Projects")

projects = [
    (
        "PodPlayer AI", "2026", True,
        "Full-stack podcast & audiobook platform — PWA + React Native app sharing one FastAPI backend.",
        ["React 18, TypeScript, React Native, Expo SDK 54, FastAPI, MongoDB, Tailwind CSS"],
        ["Audio/video streaming with background playback, lock-screen controls, and offline downloads",
         "AI playlists — 'Made For You' and Calendar Picks built from the user's upcoming events",
         "Creator Hub with episode analytics, Flash Drop recorder, and RSS podcast import",
         "In-app PodCoin store with 11 themes, avatar frames, and paywalled content"],
    ),
    (
        "SweetTeams", "2026", False,
        "Microsoft Teams-like web app — WebRTC video for 50+ participants, screen sharing & chat.",
        ["JavaScript, WebRTC, Node.js, SQLite, SendGrid, Render"],
        ["WebRTC P2P video for 50+ concurrent users",
         "Passwordless magic-link auth via SendGrid",
         "Real-time chat & screen sharing"],
    ),
    (
        "WorkSearcher", "2026", False,
        "Swedish job aggregator — React dashboard + FastAPI backend, AI-scored postings.",
        ["React, FastAPI, Python, PostgreSQL, OpenAI gpt-4o-mini, MinHash LSH, Docker"],
        ["Pluggable connector protocol: JobTech, Adzuna, Greenhouse, Lever",
         "MinHash LSH near-duplicate removal + per-job summary & relevance score (OpenAI)"],
    ),
    (
        "Song Studio", "2026", True,
        "Desktop app that turns a one-line idea into a full song — lyrics, vocals & instrumentation.",
        ["Python, OpenAI API, MusicGen, Suno Bark, PyTorch, customtkinter"],
        ["Two audio backends: local (MusicGen + Bark) or Suno cloud singing",
         "CUDA-aware: ~30 s per song on GPU"],
    ),
    (
        "VoiceClone Studio", "2026", False,
        "Clone any voice and synthesize speech — GUI + CLI, 17+ languages, zero-shot.",
        ["Python 3.11, Coqui XTTS-v2, Faster Whisper, FFmpeg, CUDA, Tkinter, PyInstaller"],
        ["Zero-shot voice cloning — no training required",
         "Standalone Windows GUI + PyInstaller single-file exe"],
    ),
    (
        "Sebbes Kokbok", "2026", False,
        "Bilingual recipe collection — 45 dishes across Swedish, Korean, Chinese & Japanese cuisines.",
        ["React, Vite, JavaScript, Node.js, Express, Docker, Nginx, Render"],
        ["Bilingual UI with full English/Swedish toggle",
         "Native script preservation: Hangul, Hanzi, Kanji/Kana"],
    ),
    (
        "SebbePyGame", "2025", True,
        "Endless wave survival ARPG — 15 spells, support gems, deep skill tree.",
        ["Python 3.11, Pygame, JSON Save"],
        ["15 auto-cast spells with 3 support-gem socket slots each",
         "Passive skill tree, 12-slot equipment, crafting, breach events"],
    ),
    (
        "Shop API Tester", "2025", False,
        "C# .NET REST API for an e-commerce shop with full CRUD.",
        ["C#, .NET, EF Core, REST, Swagger"],
        ["CRUD endpoints, Entity Framework Core migrations, Swagger/OpenAPI docs"],
    ),
]

for title, year, is_private, tagline, tech_list, highlights in projects:
    # Project title line
    p  = doc.add_paragraph()
    r1 = p.add_run(title)
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = DARK
    if is_private:
        r_priv = p.add_run("  [Private]")
        r_priv.font.size      = Pt(8)
        r_priv.font.color.rgb = LIGHT
    r2 = p.add_run(f"\t{year}")
    r2.font.size      = Pt(9)
    r2.font.color.rgb = LIGHT
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9260')
    tabs.append(tab_el)
    pPr.append(tabs)
    set_spacing(p, before=60, after=20)

    # Tagline
    tl = doc.add_paragraph(tagline)
    tl.runs[0].font.size      = Pt(9)
    tl.runs[0].font.color.rgb = MID
    tl.runs[0].italic         = True
    set_spacing(tl, before=0, after=20)

    # Tech
    tech_p = doc.add_paragraph()
    rt = tech_p.add_run("Tech: ")
    rt.bold           = True
    rt.font.size      = Pt(8.5)
    rt.font.color.rgb = ACCENT
    rv = tech_p.add_run(tech_list[0])
    rv.font.size      = Pt(8.5)
    rv.font.color.rgb = MID
    set_spacing(tech_p, before=0, after=20)

    # Highlights
    for h in highlights:
        bul = doc.add_paragraph(style='List Bullet')
        bul.text = ""
        r = bul.add_run(f"  {h}")
        r.font.size      = Pt(9)
        r.font.color.rgb = MID
        set_spacing(bul, before=0, after=10)

# ═══════════════════════════════════════════════════════════════════════════════
#  EDUCATION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Education")

edu = [
    ("2024 – 2026", "Higher Vocational Education — Software Development with AI Specialization",
     "NBI/Handelsakademin, Halmstad"),
]
for period, degree, school in edu:
    p  = doc.add_paragraph()
    r1 = p.add_run(degree)
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = DARK
    r2 = p.add_run(f"\t{period}")
    r2.font.size      = Pt(9)
    r2.font.color.rgb = LIGHT
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), '9260')
    tabs.append(tab_el)
    pPr.append(tabs)
    set_spacing(p, before=0, after=20)

    s_p = doc.add_paragraph(school)
    s_p.runs[0].font.size      = Pt(9)
    s_p.runs[0].font.color.rgb = MID
    set_spacing(s_p, before=0, after=40)

# ═══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "References")

references = [
    ("Marcus Granberg", "Manager · Podmanager.AI",
     "Supervisor during LIA internship at Podmanager.AI",
     "marcus.granberg@podmanager.ai", "076-100 04 71"),
    ("Antonio Prgomet", "Teacher · NBI/Handelsakademin",
     "Teacher and mentor during software development with AI studies",
     "antonio.prgomet@nbi-handelsakademin.se", ""),
    ("Robert Ploski", "Thesis partner & Class Mate · NBI/Handelsakademin",
     "Student colleague during software development with AI at NBI/Handelsakademin",
     "robert.ploski90@gmail.com", ""),
]

for name, role_co, relation, email, phone in references:
    p  = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold           = True
    r1.font.size      = Pt(10)
    r1.font.color.rgb = DARK
    r2 = p.add_run(f"  ·  {role_co}")
    r2.font.size      = Pt(9)
    r2.font.color.rgb = MID
    set_spacing(p, before=60, after=10)

    rel_p = doc.add_paragraph(relation)
    rel_p.runs[0].font.size      = Pt(9)
    rel_p.runs[0].font.color.rgb = MID
    set_spacing(rel_p, before=0, after=10)

    contact_parts = []
    if email: contact_parts.append(email)
    if phone: contact_parts.append(phone)
    if contact_parts:
        c_p = doc.add_paragraph("  ·  ".join(contact_parts))
        c_p.runs[0].font.size      = Pt(8.5)
        c_p.runs[0].font.color.rgb = LIGHT
        set_spacing(c_p, before=0, after=30)

# ═══════════════════════════════════════════════════════════════════════════════
#  Save
# ═══════════════════════════════════════════════════════════════════════════════
output = r"c:\Users\quo\repos\SebbePortfolio2\Sebastian_Larsson_CV.docx"
doc.save(output)
print(f"Saved: {output}")
